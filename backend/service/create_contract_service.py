from database.setup_postgres import get_db
from database.table.table_postgres import type_contract, contract, contract_history_mess, session
from fastapi import APIRouter, UploadFile, File, HTTPException
from docxtpl import DocxTemplate
from sqlalchemy import select, delete
import re
from docx import Document
from agent_chatbot.grab.grab import app_workflow
from agent_chatbot.agent_state.agent_state import ContractState
import os
from fastapi.responses import FileResponse

async def upload_file(file, db):
    """
        Hàm xử lý khi upload file template hợp đồng
        - check trùng tên file
        - check .docx
        - đọc file lấy ra json các input
        - save file và lưu thông tin lên database  
    """


    filename = file.filename
    file_db = await db.execute(select(type_contract).where(type_contract.id == filename))
    file_db = file_db.scalars().all()

    if(len(file_db)>0):
        raise HTTPException(status_code=400, detail="file đã tồn tại.")
    
    if not filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Chỉ chấp nhận file .docx")
    

    file_path = f"database/dot/template/{filename}"
    content = await file.read()
    
    with open(file_path, "wb") as f:
        f.write(content)

    print("="*50+"\n")
    result = build_vllm_json_schema(scan_template_fast(file_path))
    # print(result)
    print("\n"+"="*50)
    new_type_contract = type_contract(
        id = filename,
        path = file_path,
        required_fields = result
    )
    db.add(new_type_contract)
    await db.commit()

    return {
        "status_code":200,
        "description": "ok"
    }


async def create_contract(create_contract_request, db):
    """
        Hàm tạo một hợp đồng theo yêu cầu của người dùng với chatbot
        - gọi qua workflow của agent để  nó tự trao đổi với user và thu thập thông tin từ inpur
    """


    # lấy ra template đang muốn tương tác
    stmt_template = select(type_contract).where(
        type_contract.id == create_contract_request.type_id
    )
    result_template = await db.execute(stmt_template)
    json_data_contract = result_template.scalars().first()  


    # Lấy ra session đang chat nếu không thấy thì tạo
    if(create_contract_request.session_id==-1):
        session_db = session(
            id_user = create_contract_request.user_id
        )
        db.add(session_db)
        await db.commit()
        await db.refresh(session_db)
    else:
        session_db = await db.execute(select(session).where(session.id==create_contract_request.session_id))
        session_db = session_db.scalars().first()

    # Lấy ra input của hợp đồng đang làm
    stmt_history = select(contract).where(
        contract.id_user == create_contract_request.user_id,
        contract.session == session_db.id,
        contract.id_type == create_contract_request.type_id,
        contract.status != "approve"
    )
    result_history = await db.execute(stmt_history)
    contract_history = result_history.scalars().first()

    data_input_null = None
    json_data = None
    check_save = False
    if(contract_history!=None):
        data_input_null = contract_history.missing_fields
        json_data = contract_history.json_data
        print(f"- Làm tiếp hợp đồng: {data_input_null}")
    else:
        data_input_null = ["Không có dữ liệu cũ"]
        json_data = []
        check_save = True
        print("- Không có hợp đồng nào đang làm")

    state = {
        "user_id": "1",
        "session_id": session_db.id,
        "template_id":create_contract_request.type_id, #Mã của word mău
        "target_schema":json_data_contract.required_fields, #Json mẫu của template
        "user_input": create_contract_request.user_input, #query người dùng
        "data_history_input": json_data,  
        "data_input_null": data_input_null,
        "mess":[],
        "status":"pending",
    }
    # print(state)

    result = await app_workflow.ainvoke(state)
    print("="*100+"\n")
    print("- Xong")
    print("\n"+"="*100)
    # print(state.get('data_input_null'))
    # print(state.get('data_history_input'))
    # print(state.get('status'))
    # print(result)
    if(check_save):
        # Lưu khi đây là cho hợp đồng mới
        print("- Đã tạo hợp đồng mới")
        contract_insert = contract(
            id_type = result.get('template_id'),
            id_user = result.get('user_id'),
            session = result.get('session_id'),
            status = result.get('status'),
            missing_fields = result.get('data_input_null'),
            json_data = result.get('data_history_input'),
            file_url = ""
        )
        db.add(contract_insert)
        await db.commit()
    else:
        # Cập nhật vì đây là hợp đồng cũ
        print("- Đã cập nhật lại hợp đồng")
        contract_history.missing_fields = result.get('data_input_null')
        contract_history.json_data = result.get('data_history_input')
        contract_history.status = result.get('status')
        await db.commit()

    if(result.get('status')=="approve"):
        file_path = f"database/dot/contract/{result.get('template_id')}"
        contract_history_mess_user = contract_history_mess(
            id_user = result.get("user_id"),
            session = result.get("session_id"),
            mess = result.get("user_input"),
            role="user"
        )
        contract_history_mess_chatbot = contract_history_mess(
            id_user = result.get("user_id"),
            session = result.get("session_id"),
            mess = "Đã tạo xong file",
            role="chatbot"
        )


        db.add(contract_history_mess_user)
        db.add(contract_history_mess_chatbot)
        await db.commit()

        # print(file_path)
        
        return{
            "status":200,
            "check":True,
            "session":result.get("session_id"),
            "result":"http://localhost/api/v1/contracts/download/"+result.get("template_id")
        }
    else:
        contract_history_mess_user = contract_history_mess(
            id_user = result.get("user_id"),
            session = result.get("session_id"),
            mess = result.get("user_input"),
            role="user"
        )
        contract_history_mess_chatbot = contract_history_mess(
            id_user = result.get("user_id"),
            session = result.get("session_id"),
            mess = result.get("mess"),
            role="chatbot"
        )


        db.add(contract_history_mess_user)
        db.add(contract_history_mess_chatbot)
        await db.commit()

        return {
            "status":200,
            "check":False,
            "session":result.get("session_id"),
            "mess":result.get("mess")
        }



def download(id):
    file_path = "database/dot/contract/"+id
    
    # 2. Kiểm tra file có tồn tại không
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File không tồn tại")
    
    # 3. Trả về file
    return FileResponse(
        path=file_path,
        filename=id,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )



async def load_history(request_history, db):
    
    query = select(contract_history_mess
    ).where(
        contract_history_mess.session == request_history.session_id,
        contract_history_mess.id_user == request_history.user_id
    ).order_by(
        contract_history_mess.id.asc() # Ưu tiên 2: Trong cùng status, cái nào mới hơn thì lên trên
    )

    data = await db.execute(query)
    data = data.scalars().all()



    return {
        "status":200,
        "result":data
    }

async def load_session_service(db):
    query  = await db.execute(select(session.id).order_by(session.id.asc()))
    query = query.scalars().all()

    return {
        "status":200,
        "result":query
    }

async def delete_session_service(id, db):
    query_delete_contract = await db.execute(delete(contract).where(contract.session==id))
    query_delete_history = await db.execute(delete(contract_history_mess).where(contract_history_mess.session==id))
    query_delete_session  = await db.execute(delete(session).where(session.id == id))
    await db.commit()

    return {
        "status":200,
        "result":"ok"
    }





def scan_template_fast(file_path):
    doc = Document(file_path)
    static_pattern = r'\{\{\s*((?!p\.)[a-zA-Z0-9_.]+)\s*\}\}' 
    loop_pattern = r'\{%\s*for\s+\w+\s+in\s+([a-zA-Z0-9_]+)\s*%\}'
    col_pattern = r'\{\{\s*p\.([a-zA-Z0-9_]+)\s*\}\}'
    
    all_static_fields = set()
    table_configs = {}
    full_text = ""
    
    for para in doc.paragraphs:
        full_text += para.text + " "
    
    raw_statics = re.findall(static_pattern, full_text)
    all_static_fields.update([f for f in raw_statics if "__" not in f])

    for table in doc.tables:
        current_table_key = None
        current_columns = set()
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                loop_match = re.search(loop_pattern, text)
                if loop_match:
                    key = loop_match.group(1)
                    if "__" not in key:
                        current_table_key = key
                cols = re.findall(col_pattern, text)
                current_columns.update([c for c in cols if "__" not in c])
                statics = re.findall(static_pattern, text)
                all_static_fields.update([s for s in statics if "__" not in s])
        
        if current_table_key:
            table_configs[current_table_key] = list(current_columns)

    return {
        "static_fields": list(all_static_fields),
        "tables": [
            {"table_key": k, "columns": v} for k, v in table_configs.items()
        ]
    }



def build_vllm_json_schema(template_info):
    """
        Chuyển đổi thông tin từ file Word sang JSON Schema để ép kiểu vLLM
    """
    properties = {}
    
    for field in template_info["static_fields"]:
        if any(word in field for word in ['tong', 'gia', 'so_luong', 'vat', 'tien','phan_tram_thue']):
            properties[field] = {"type": "integer"}
        else:
            properties[field] = {"type": "string"}
            
    for table in template_info["tables"]:
        table_key = table["table_key"]
        columns = table["columns"]
        
        item_properties = {}
        for col in columns:
            if any(word in col for word in ['so_luong', 'gia', 'tien', 'phan_tram_thue']):
                item_properties[col] = {"type": "integer"}
            else:
                item_properties[col] = {"type": "string"}
        
        properties[table_key] = {
            "type": "array",
            "items": {
                "type": "object",
                "properties": item_properties
            }
        }
        
    return {
        "type": "object",
        "properties": properties
    }



async def load_template_name(db):
    contract_type_ = await db.execute(select(type_contract.id))
    contract_type_ = contract_type_.scalars().all()

    return {
        "status":200,
        "result":contract_type_
    }  


async def delete_template_service(db, id_template):
    

   
    query = delete(contract).where(contract.id_type==id_template)
    await db.execute(query)
    await db.commit()

    query = delete(type_contract).where(type_contract.id == id_template)
    await db.execute(query)
    await db.commit()

    return {
        "status":200,
        "result":"ok"
    }